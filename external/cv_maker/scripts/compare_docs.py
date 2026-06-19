# Copyright 2026 Justin Cook
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_alignment_str(align_val):
    if align_val == WD_PARAGRAPH_ALIGNMENT.CENTER: return "CENTER"
    if align_val == WD_PARAGRAPH_ALIGNMENT.RIGHT: return "RIGHT"
    if align_val == WD_PARAGRAPH_ALIGNMENT.LEFT: return "LEFT"
    if align_val == WD_PARAGRAPH_ALIGNMENT.JUSTIFY: return "JUSTIFY"
    return "NONE (Inherit)"

def analyze(path, label):
    print(f"\n--- {label} ({path}) ---")
    doc = Document(path)
    
    # 1. Page Margins (Section 0)
    sect = doc.sections[0]
    print(f"Margins: L={sect.left_margin}, R={sect.right_margin}, T={sect.top_margin}")
    
    # 2. First 5 Paragraphs (Header area)
    print("Header Area (First 5 paragraphs):")
    for i, p in enumerate(doc.paragraphs[:5]):
        text = p.text.strip()
        if not text: continue
        align = get_alignment_str(p.alignment)
        print(f"  [{i}] Style='{p.style.name}' Align='{align}' Text='{text[:20]}...'")

    # 3. Find 'Experience' or Main Headers
    print("Section Headers:")
    for p in doc.paragraphs:
        if 'Experience' in p.text or 'Education' in p.text:
             align = get_alignment_str(p.alignment)
             print(f"  Header Found: Style='{p.style.name}' Align='{align}' Text='{p.text[:30]}...'")

if __name__ == "__main__":
    analyze("Justin_Cook_Engineering_Lead.docx", "TEMPLATE")
    analyze("Tailored_CV.docx", "GENERATED")
