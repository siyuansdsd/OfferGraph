#!/usr/bin/env python3
# Copyright 2026 Justin Cook
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""
My CV Template
pip install python-docx
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_cv():
    document = Document()

    # Style configuration
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- HEADER ---
    head = document.add_heading('[Your Name]', 0)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    contact = document.add_paragraph('[Your Job Title / Headline]')
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.runs[0].bold = True
    
    details = document.add_paragraph('[City, Country] | [Phone Number] | [Email Address]\ngithub.com/[username]')
    details.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()  # Spacer

    # --- EXECUTIVE SUMMARY ---
    document.add_heading('EXECUTIVE SUMMARY', level=1)
    summary = document.add_paragraph(
        '[Insert your executive summary here. Describe your experience, key skills, and what you bring to the table.]'
    )

    # --- CORE COMPETENCIES ---
    document.add_heading('CORE COMPETENCIES', level=1)
    competencies = [
        ('Category 1:', 'Skill A, Skill B, Skill C.'),
        ('Category 2:', 'Skill D, Skill E, Skill F.'),
        ('Category 3:', 'Skill G, Skill H, Skill I.')
    ]
    
    for category, skills in competencies:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(category).bold = True
        p.add_run(f" {skills}")

    # --- PROFESSIONAL EXPERIENCE ---
    document.add_heading('PROFESSIONAL EXPERIENCE', level=1)

    # Latest Job
    p = document.add_paragraph()
    p.add_run('[LATEST COMPANY]').bold = True
    p.add_run(' | [Location] | ')
    p.add_run('[Dates]').italic = True
    
    p = document.add_paragraph()
    p.add_run('[Job Title]').bold = True
    document.add_paragraph('[Brief summary of your role and impact.]').italic = True
    
    bullets = [
        ('Achievement 1:', 'Description of achievement.'),
        ('Achievement 2:', 'Description of achievement.'),
        ('Achievement 3:', 'Description of achievement.')
    ]
    for title, desc in bullets:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    document.add_paragraph() # Spacer

    # Previous Job
    p = document.add_paragraph()
    p.add_run('[PREVIOUS COMPANY]').bold = True
    p.add_run(' | [Location] | ')
    p.add_run('[Dates]').italic = True
    
    p = document.add_paragraph()
    p.add_run('[Job Title]').bold = True
    document.add_paragraph('[Brief summary of your role and impact.]').italic = True
    
    bullets_prev = [
        ('Achievement 1:', 'Description of achievement.'),
        ('Achievement 2:', 'Description of achievement.')
    ]
    for title, desc in bullets_prev:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    document.add_paragraph() # Spacer

    document.add_paragraph('Prior roles include [Role] at [Company], and [Role] at [Company].').italic = True

    # --- PROJECTS ---
    document.add_heading('TECHNICAL PROJECTS & OPEN SOURCE', level=1)
    document.add_paragraph('Visible at: github.com/[username]').italic = True

    projects = [
        ('Project A:', 'Description of project.'),
        ('Project B:', 'Description of project.')
    ]
    for title, desc in projects:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    # --- EDUCATION ---
    document.add_heading('EDUCATION & CERTIFICATIONS', level=1)
    document.add_paragraph('[Degree], [University]')
    
    certs = document.add_paragraph()
    certs.add_run('Certifications: ').bold = True
    certs.add_run('[Certification 1] | [Certification 2]')

    # Save
    document.save('My_CV.docx')
    print("CV generated successfully: My_CV.docx")

if __name__ == "__main__":
    create_cv()
