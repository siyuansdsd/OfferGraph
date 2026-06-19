# Copyright 2026 Justin Cook
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import unittest
from docx import Document
from cv_maker.generator import CVGenerator

class TestCVGeneratorAssembly(unittest.TestCase):
    
    def setUp(self):
        # Create a real in-memory Document for testing logic
        self.doc = Document()
        # Create generator with this document
        self.generator = CVGenerator()
        self.generator.document = self.doc
        # Reset buffers
        self.generator.header_elements = []
        self.generator.injections = {}
        # Clear default paragraph if it exists to ensure clean state
        if len(self.doc.element.body) > 0:
             self.doc.element.body.clear()
        
        print(f"\nDEBUG: Initial body length: {len(self.doc.element.body)}")

    def test_prepend_content_buffers(self):
        """Verify _prepend_content adds to header_elements and clears body."""
        
        # Define content function that adds paragraphs
        def content_func():
            self.doc.add_paragraph("Header Content 1")
            self.doc.add_paragraph("Header Content 2")
            
        # Initial state: empty body
        self.assertEqual(len(self.doc.element.body), 0)
        
        # Run prepend
        self.generator._prepend_content(content_func)
        
        # Verify body is still empty (content moved to buffer)
        # Note: Depending on implementation, body might be empty or have 0 elements
        # The key is that the NEW elements are gone from body
        self.assertEqual(len(self.doc.element.body), 0)
        
        # Verify buffer has 2 elements
        self.assertEqual(len(self.generator.header_elements), 2)
        
    def test_inject_content_after_buffers(self):
        """Verify _inject_content_after buffers content mapped to target."""
        
        # Setup: Add a target paragraph to body
        target_p = self.doc.add_paragraph("Target Section")
        
        # Content function
        def content_func():
            self.doc.add_paragraph("Injected 1")
            self.doc.add_paragraph("Injected 2")
            
        # Run inject
        self.generator._inject_content_after(target_p, content_func)
        
        # Verify body only has target (injected elements removed)
        self.assertEqual(len(self.doc.element.body), 1) # Just the target
        self.assertEqual(self.doc.paragraphs[0].text, "Target Section")
        
        # Verify buffer has elements mapped to target element
        target_elem = target_p._element
        self.assertIn(target_elem, self.generator.injections)
        self.assertEqual(len(self.generator.injections[target_elem]), 2)

    def test_assemble_document_order(self):
        """Verify _assemble_document reconstructs body in correct order."""
        
        # 1. Setup Header Elements (Mocked from previous step)
        # We manually create elements and add to buffer
        self.doc.add_paragraph("HEADER")
        header_elem = self.doc.paragraphs[0]._element
        # Remove from body to simulate buffering
        self.doc.element.body.remove(header_elem)
        self.generator.header_elements = [header_elem]
        
        # 2. Setup Body with Target
        target_p = self.doc.add_paragraph("TEMPLATE SECTION")
        target_elem = target_p._element
        
        # 3. Setup Injection for Target
        # Add injected content temporarily to get elements
        self.doc.add_paragraph("INJECTED CONTENT")
        injected_elem = self.doc.paragraphs[1]._element # Index 1 because target is 0
        self.doc.element.body.remove(injected_elem)
        
        self.generator.injections = {target_elem: [injected_elem]}
        
        # Check pre-assembly state: Body has 1 element (Target)
        self.assertEqual(len(self.doc.element.body), 1)
        
        # 4. Run Assemble
        self.generator._assemble_document()
        
        # 5. Verify Order: Header -> Target -> Injected
        body_elems = self.doc.element.body
        self.assertEqual(len(body_elems), 3)
        
        # Note: python-docx wraps elements in proxy objects, but we can check underlying xml or identity
        # Re-fetch paragraphs to check text
        paragraphs = self.doc.paragraphs
        self.assertEqual(paragraphs[0].text, "HEADER")
        self.assertEqual(paragraphs[1].text, "TEMPLATE SECTION")
        self.assertEqual(paragraphs[2].text, "INJECTED CONTENT")
    def test_generate_section_order(self):
        """Verify generate() places sections in correct order and keeps earlier roles inside experience."""
        from cv_maker.models import CVData, Experience, EarlierExperience
        import tempfile, os

        gen = CVGenerator()

        data = CVData(
            name="Test User",
            title="Engineer",
            contact_info="test@example.com",
            executive_summary="Summary text.",
            competencies=[("Category:", "Skill 1, Skill 2.")],
            experience=[
                Experience(
                    title="Senior Engineer",
                    company="ACME",
                    location="City",
                    dates="Jan 2022 – Present",
                    bullets=[("Delivery:", "Led project X.")]
                )
            ],
            earlier_experience=[
                EarlierExperience(
                    title="Junior Engineer",
                    company="Old Corp",
                    summary="Did early career work."
                )
            ],
            projects=[("Proj:", "Description.")],
            education=["BS Computer Science"],
            certifications="Cert A"
        )

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            gen.generate(data, output_path)

            from docx import Document
            doc = Document(output_path)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]

            # Find heading positions
            heading_indices = {}
            for i, text in enumerate(texts):
                if text == 'EXECUTIVE SUMMARY':
                    heading_indices['summary'] = i
                elif text == 'CORE COMPETENCIES':
                    heading_indices['competencies'] = i
                elif text == 'PROFESSIONAL EXPERIENCE':
                    heading_indices['experience'] = i

            # Assert correct order
            self.assertIn('summary', heading_indices, "EXECUTIVE SUMMARY heading not found")
            self.assertIn('competencies', heading_indices, "CORE COMPETENCIES heading not found")
            self.assertIn('experience', heading_indices, "PROFESSIONAL EXPERIENCE heading not found")
            self.assertNotIn('EARLIER CAREER EXPERIENCE', texts)
            self.assertIn('Junior Engineer, Old Corp', texts)

            self.assertLess(heading_indices['summary'], heading_indices['competencies'],
                           "Summary must come before Competencies")
            self.assertLess(heading_indices['competencies'], heading_indices['experience'],
                           "Competencies must come before Experience")
            self.assertLess(heading_indices['experience'], texts.index('Junior Engineer, Old Corp'),
                           "Earlier roles must appear inside Professional Experience after detailed roles")
        finally:
            os.unlink(output_path)

    def test_experience_preserves_input_order(self):
        """Verify generator preserves the order of experience entries as given (LLM contract)."""
        from cv_maker.models import CVData, Experience
        import tempfile, os

        gen = CVGenerator()

        data = CVData(
            name="Test User",
            title="Engineer",
            contact_info="test@example.com",
            executive_summary="Summary.",
            competencies=[],
            experience=[
                Experience(title="Role A", company="COMPANY A", location="City", dates="2023 – Present", bullets=[]),
                Experience(title="Role B", company="COMPANY B", location="City", dates="2020 – 2023", bullets=[]),
                Experience(title="Role C", company="COMPANY C", location="City", dates="2017 – 2020", bullets=[]),
            ]
        )

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            gen.generate(data, output_path)

            from docx import Document
            doc = Document(output_path)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]

            # Find company lines (they appear as "COMPANY X | City | dates")
            company_lines = [t for t in texts if t.startswith('COMPANY')]
            self.assertEqual(len(company_lines), 3)
            self.assertTrue(company_lines[0].startswith('COMPANY A'))
            self.assertTrue(company_lines[1].startswith('COMPANY B'))
            self.assertTrue(company_lines[2].startswith('COMPANY C'))
        finally:
            os.unlink(output_path)

if __name__ == '__main__':
    unittest.main()
